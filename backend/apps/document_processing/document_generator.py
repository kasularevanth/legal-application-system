# apps/document_processing/document_generator.py
import os
import tempfile
import logging
from typing import Dict, List, Optional
from datetime import datetime
import uuid

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from azure.storage.blob import BlobServiceClient, BlobClient
from azure.core.exceptions import AzureError

from django.conf import settings
from django.utils import timezone
from ..models import DocumentTemplate, LegalCase

logger = logging.getLogger(__name__)

class AzureBlobManager:
    """Manage Azure Blob Storage operations"""
    
    def __init__(self):
        self.account_name = getattr(settings, 'AZURE_ACCOUNT_NAME', None)
        self.account_key = getattr(settings, 'AZURE_ACCOUNT_KEY', None)
        self.container_name = getattr(settings, 'AZURE_CONTAINER', 'legal-documents')
        
        if self.account_name and self.account_key:
            self.blob_service_client = BlobServiceClient(
                account_url=f"https://{self.account_name}.blob.core.windows.net",
                credential=self.account_key
            )
        else:
            logger.warning("Azure credentials not configured, using local storage")
            self.blob_service_client = None

    def upload_template(self, file_path: str, blob_name: str) -> str:
        """Upload a template file to Azure Blob Storage"""
        try:
            if not self.blob_service_client:
                return f"file://{file_path}"  # Local fallback
            
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            
            with open(file_path, 'rb') as data:
                blob_client.upload_blob(data, overwrite=True)
            
            return blob_client.url
            
        except AzureError as e:
            logger.error(f"Failed to upload template to Azure: {e}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error uploading template: {e}")
            raise

    def download_template(self, blob_name: str, local_path: str) -> str:
        """Download a template from Azure Blob Storage"""
        try:
            if not self.blob_service_client:
                # For local development, assume templates are in media/templates/
                template_path = os.path.join(settings.MEDIA_ROOT, 'templates', blob_name)
                if os.path.exists(template_path):
                    return template_path
                raise FileNotFoundError(f"Template not found: {template_path}")
            
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            
            with open(local_path, 'wb') as download_file:
                download_file.write(blob_client.download_blob().readall())
            
            return local_path
            
        except AzureError as e:
            logger.error(f"Failed to download template from Azure: {e}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error downloading template: {e}")
            raise

    def upload_generated_document(self, file_path: str, blob_name: str) -> str:
        """Upload generated document to Azure Blob Storage"""
        try:
            if not self.blob_service_client:
                # For local development, move to media/documents/
                docs_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
                os.makedirs(docs_dir, exist_ok=True)
                local_path = os.path.join(docs_dir, blob_name)
                
                import shutil
                shutil.copy2(file_path, local_path)
                return f"{settings.MEDIA_URL}documents/{blob_name}"
            
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=f"generated/{blob_name}"
            )
            
            with open(file_path, 'rb') as data:
                blob_client.upload_blob(data, overwrite=True)
            
            return blob_client.url
            
        except AzureError as e:
            logger.error(f"Failed to upload document to Azure: {e}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error uploading document: {e}")
            raise


class WordDocumentGenerator:
    """Generate Word documents from templates"""
    
    def __init__(self):
        self.blob_manager = AzureBlobManager()

    def generate_document(self, case: LegalCase) -> Dict:
        """Generate a Word document for the legal case"""
        try:
            # Get the appropriate template
            template = self._get_template_for_case(case)
            if not template:
                raise ValueError("No template found for case type")

            # Download template from Azure Blob
            with tempfile.TemporaryDirectory() as temp_dir:
                template_path = os.path.join(temp_dir, f"template_{template.id}.docx")
                self.blob_manager.download_template(template.blob_name, template_path)

                # Load the template document
                doc = Document(template_path)

                # Replace placeholders with actual values
                replacements = self._prepare_replacements(case, template)
                self._replace_placeholders(doc, replacements)

                # Generate output filename
                output_filename = f"legal_doc_{case.case_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                output_path = os.path.join(temp_dir, output_filename)

                # Save the generated document
                doc.save(output_path)

                # Upload to Azure Blob
                document_url = self.blob_manager.upload_generated_document(
                    output_path, output_filename
                )

                # Update case with document information
                case.generated_document_url = document_url
                case.document_blob_name = output_filename
                case.template_used = template
                case.status = 'document_ready'
                case.completed_at = timezone.now()
                case.save()

                case.add_processing_step("Document generated successfully")

                return {
                    'success': True,
                    'document_url': document_url,
                    'document_id': str(case.case_id),
                    'template_used': template.name,
                    'blob_name': output_filename,
                    'preview_text': self._generate_preview(doc)
                }

        except Exception as e:
            logger.error(f"Document generation failed for case {case.case_id}: {e}")
            case.status = 'error'
            case.error_details = f"Document generation failed: {str(e)}"
            case.save()
            case.add_processing_step(f"Document generation failed: {str(e)}")
            
            return {
                'success': False,
                'error': str(e),
                'case_id': str(case.case_id)
            }

    def _get_template_for_case(self, case: LegalCase) -> Optional[DocumentTemplate]:
        """Get the appropriate template for the case type"""
        if not case.detected_case_type:
            return None
        
        # Get the most recent active template for this case type
        template = DocumentTemplate.objects.filter(
            case_type_mapping=case.detected_case_type,
            is_active=True
        ).order_by('-created_at').first()
        
        return template

    def _prepare_replacements(self, case: LegalCase, template: DocumentTemplate) -> Dict[str, str]:
        """Prepare replacement values for template placeholders"""
        replacements = {}
        
        # Basic case information
        replacements.update({
            '${case_id}': str(case.case_id),
            '${date_created}': case.created_at.strftime('%B %d, %Y'),
            '${case_type}': case.detected_case_type.case_type if case.detected_case_type else '',
            '${user_name}': f"{case.user.first_name} {case.user.last_name}".strip() or case.user.username,
            '${user_email}': case.user.email,
            '${user_phone}': case.user.phone_number or '',
        })

        # Map answers to template fields using field mappings
        if template.field_mappings and case.answers_received:
            for template_field, question_field in template.field_mappings.items():
                # Find the answer for this field
                for question, answer in case.answers_received.items():
                    # Match by question text or field name
                    if question_field in question or question_field == question:
                        replacements[f'${{{template_field}}}'] = answer
                        break

        # Add any missing fields with empty values
        for field in template.field_mappings.keys():
            placeholder = f'${{{field}}}'
            if placeholder not in replacements:
                replacements[placeholder] = '[TO BE FILLED]'

        # Add current date/time
        replacements.update({
            '${current_date}': datetime.now().strftime('%B %d, %Y'),
            '${current_time}': datetime.now().strftime('%I:%M %p'),
            '${document_generated_at}': datetime.now().strftime('%B %d, %Y at %I:%M %p'),
        })

        logger.info(f"Prepared {len(replacements)} replacements for case {case.case_id}")
        return replacements

    def _replace_placeholders(self, doc: Document, replacements: Dict[str, str]):
        """Replace placeholders in the Word document"""
        try:
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        # Replace maintaining formatting
                        self._replace_text_in_paragraph(paragraph, placeholder, value)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    self._replace_text_in_paragraph(paragraph, placeholder, value)

            # Replace in headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, value)
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, value)

        except Exception as e:
            logger.error(f"Failed to replace placeholders: {e}")
            raise

    def _replace_text_in_paragraph(self, paragraph, placeholder: str, value: str):
        """Replace text in a paragraph while preserving formatting"""
        try:
            if placeholder in paragraph.text:
                # Find runs that contain the placeholder
                runs = paragraph.runs
                for i, run in enumerate(runs):
                    if placeholder in run.text:
                        # Simple replacement - could be enhanced to preserve formatting better
                        run.text = run.text.replace(placeholder, value)
                        break
                else:
                    # Placeholder spans multiple runs, use simple text replacement
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    
        except Exception as e:
            logger.warning(f"Failed to replace text in paragraph: {e}")
            # Fallback to simple replacement
            paragraph.text = paragraph.text.replace(placeholder, value)

    def _generate_preview(self, doc: Document) -> str:
        """Generate a text preview of the document"""
        try:
            preview_text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    preview_text += paragraph.text.strip() + "\n\n"
                    paragraph_count += 1
                    
                    # Limit preview to first 5 paragraphs
                    if paragraph_count >= 5:
                        break
            
            if len(preview_text) > 500:
                preview_text = preview_text[:500] + "..."
            
            return preview_text
            
        except Exception as e:
            logger.error(f"Failed to generate preview: {e}")
            return "Preview not available"

    def create_template(self, case_type_mapping, template_name: str, template_file_path: str, field_mappings: Dict) -> DocumentTemplate:
        """Create a new document template"""
        try:
            # Generate blob name
            blob_name = f"templates/{case_type_mapping.uuid}_{template_name.replace(' ', '_')}.docx"
            
            # Upload template to Azure Blob
            blob_url = self.blob_manager.upload_template(template_file_path, blob_name)
            
            # Create template record
            template = DocumentTemplate.objects.create(
                case_type_mapping=case_type_mapping,
                name=template_name,
                blob_url=blob_url,
                blob_container=self.blob_manager.container_name,
                blob_name=blob_name,
                field_mappings=field_mappings,
                is_active=True
            )
            
            logger.info(f"Created template {template.id} for case type {case_type_mapping.case_type}")
            return template
            
        except Exception as e:
            logger.error(f"Failed to create template: {e}")
            raise

    def preview_document(self, case: LegalCase) -> Dict:
        """Generate a preview of the document without saving"""
        try:
            template = self._get_template_for_case(case)
            if not template:
                return {'error': 'No template found for case type'}

            with tempfile.TemporaryDirectory() as temp_dir:
                template_path = os.path.join(temp_dir, f"template_{template.id}.docx")
                self.blob_manager.download_template(template.blob_name, template_path)

                doc = Document(template_path)
                replacements = self._prepare_replacements(case, template)
                self._replace_placeholders(doc, replacements)

                preview_text = self._generate_preview(doc)
                
                return {
                    'success': True,
                    'preview_text': preview_text,
                    'field_mappings': template.field_mappings,
                    'missing_fields': self._find_missing_fields(case, template)
                }

        except Exception as e:
            logger.error(f"Preview generation failed: {e}")
            return {'error': str(e)}

    def _find_missing_fields(self, case: LegalCase, template: DocumentTemplate) -> List[str]:
        """Find fields that are required but not filled"""
        missing = []
        
        if template.field_mappings and case.answers_received:
            for template_field, question_field in template.field_mappings.items():
                found = False
                for question, answer in case.answers_received.items():
                    if question_field in question and answer.strip():
                        found = True
                        break
                
                if not found:
                    missing.append(template_field)
        
        return missing